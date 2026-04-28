import streamlit as st
import pandas as pd
import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io


# --- HELPER FUNCTIONS ---
def get_malay_date(date_obj):
    months = ["Januari", "Februari", "Mac", "April", "Mei", "Jun", 
              "Julai", "Ogos", "September", "Oktober", "November", "Disember"]
    days = ["Isnin", "Selasa", "Rabu", "Khamis", "Jumaat", "Sabtu", "Ahad"]
    day_name = days[date_obj.weekday()]
    month_name = months[date_obj.month - 1]
    return f"{date_obj.day} {month_name} {date_obj.year} ({day_name})"


def get_epid_week(date_obj):
    # Epid Week 1 starts on 04/01/2026
    start_date = datetime.date(2026, 1, 4)
    if date_obj < start_date:
        return 52 # Fallback for dates before 2026
    delta = (date_obj - start_date).days
    return (delta // 7) + 1


def set_cell_background(cell, color_hex):
    # Helper to set table header background color
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading_elm)


# --- MAIN APP ---
st.set_page_config(page_title="Generator Laporan BWKK CPRC Selangor", layout="wide")
st.title("Generator Laporan Harian BWKK CPRC Selangor")
st.write("Sila muat naik fail e-Notifikasi dan Senarai Wabak (Outbreak Listing) untuk menjana laporan.")


# Upload files
file_enotif = st.file_uploader("Muat Naik Fail E-Notifikasi (Excel/CSV)", type=['xlsx', 'xls', 'csv'])
file_outbreak = st.file_uploader("Muat Naik Fail Senarai Wabak/Outbreak Listing (Excel/CSV)", type=['xlsx', 'xls', 'csv'])


# Only show generate button if both files are uploaded
if file_enotif and file_outbreak:
    if st.button("Jana Laporan (Generate Report)"):
        with st.spinner("Sedang memproses data dan menjana laporan..."):
            try:
               # 1. READ UPLOADED FILES
                if file_enotif.name.endswith('csv'):
                    df_enotif = pd.read_csv(file_enotif)
                else:
                    df_enotif = pd.read_excel(file_enotif)
                    
                if file_outbreak.name.endswith('csv'):
                    df_outbreak = pd.read_csv(file_outbreak)
                else:
                    df_outbreak = pd.read_excel(file_outbreak)

                # --- NEW FIX: Clean up column names to prevent 'KeyError' ---
                # Remove any invisible spaces from the headers in both files
                df_enotif.columns = df_enotif.columns.astype(str).str.strip()
                df_outbreak.columns = df_outbreak.columns.astype(str).str.strip()
                
                # Automatically rename the disease column to exactly 'Penyakit' regardless of its original case
                for col in df_enotif.columns:
                    if col.upper() == 'PENYAKIT':
                        df_enotif.rename(columns={col: 'Penyakit'}, inplace=True)
                        break

                # 2. DATE CALCULATIONS
                today = datetime.date.today()
                yesterday = today - datetime.timedelta(days=1)
                
                tarikh_today_str = get_malay_date(today)
                tarikh_yesterday_str = get_malay_date(yesterday)
                epid_week = get_epid_week(today)


                # 3. PROCESS DATA
                # SECTION 1.0: E-Notifikasi
                total_enotif = len(df_enotif)
                # Assuming 'Penyakit' and column 'BQ' are in the dataset
                col_bq_name = df_enotif.columns if len(df_enotif.columns) > 68 else df_enotif.columns[-1] # Fallback to last col if BQ doesn't exist exactly
                df_jadual_1 = df_enotif.groupby(['Penyakit', col_bq_name]).size().unstack(fill_value=0)
                if 'Average' in df_jadual_1.columns:
                    df_jadual_1 = df_jadual_1.drop(columns=['Average'])
                df_jadual_1['JUMLAH'] = df_jadual_1.sum(axis=1)


                # SECTION 2.0: Notifikasi Wabak
                # Filter outbreaks from 04/01/2026 til yesterday
                # Assuming 'F' is the 5th index, 'AL' is the 37th index
                col_penyakit_outbreak = df_outbreak.columns 
                col_date_isytihar = df_outbreak.columns 
                
                df_outbreak[col_date_isytihar] = pd.to_datetime(df_outbreak[col_date_isytihar], errors='coerce').dt.date
                mask = (df_outbreak[col_date_isytihar] >= datetime.date(2026, 1, 4)) & (df_outbreak[col_date_isytihar] <= yesterday)
                df_outbreak_filtered = df_outbreak[mask]
                
                total_outbreak = len(df_outbreak_filtered[df_outbreak_filtered[col_date_isytihar] == yesterday])
                
                # Group for Jadual 2
                df_jadual_2 = df_outbreak_filtered.groupby(col_penyakit_outbreak).apply(
                    lambda x: pd.Series({
                        'HARIAN': (x[col_date_isytihar] == yesterday).sum(),
                        'KUMULATIF': len(x)
                    })
                ).reset_index()
                df_jadual_2.rename(columns={col_penyakit_outbreak: 'PENYAKIT'}, inplace=True)
                
                if df_jadual_2.empty:
                    df_jadual_2 = pd.DataFrame(columns=['PENYAKIT', 'HARIAN', 'KUMULATIF'])


                # SECTION 3.0: Wabak Vektor (Fetching via public Google Sheet)
                sheet_vektor_id = "1bjyNcntm-I6nRaIVkVdJqJRAzn5r2tYFfjUAN0emv9w" 
                gid_vektor = "0"
                url_vektor = f"https://docs.google.com/spreadsheets/d/{sheet_vektor_id}/export?format=csv&gid={gid_vektor}&range=N21:T32"
                try:
                    df_jadual_3 = pd.read_csv(url_vektor)
                except:
                    # Fallback empty dataframe if sheet fails to load
                    df_jadual_3 = pd.DataFrame(columns=["DAERAH", "DENGGI HARIAN", "DENGGI KUM", "MALARIA HARIAN", "MALARIA KUM", "CHIKUNGUNYA HARIAN", "CHIKUNGUNYA KUM"])


                # SECTION 4.0: Kejadian Insiden (Fetching via public Google Sheet)
                sheet_bencana_id = "1Fp6IORRfdWSJCTC8vqSSoQz6RpCpNXHzO6jj0tHEf2c"
                gid_bencana_2026 = "1342717767"
                gid_bencana_table = "1342717767"
                
                url_bencana_2026 = f"https://docs.google.com/spreadsheets/d/{sheet_bencana_id}/export?format=csv&gid={gid_bencana_2026}"
                url_bencana_table = f"https://docs.google.com/spreadsheets/d/{sheet_bencana_id}/export?format=csv&gid={gid_bencana_table}&range=AH2:AU13"
                
                try:
                    df_bencana_2026 = pd.read_csv(url_bencana_2026)
                    df_jadual_4 = pd.read_csv(url_bencana_table)
                    # Check column C (index 2) for yesterday's date
                    col_c_dates = pd.to_datetime(df_bencana_2026.iloc[:, 2], format='%d/%m/%Y', errors='coerce').dt.date
                    total_insiden_yesterday = (col_c_dates == yesterday).sum()
                except:
                    total_insiden_yesterday = 0
                    df_jadual_4 = pd.DataFrame()


                # --- 4. DOCX GENERATION ---
                doc = Document()
                
                # Add Logo (Make sure 'logo_kkm.png' is in the directory)
                try:
                    doc.add_picture('logo_kkm.png', width=Inches(1.5))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                except:
                    pass # Skip if logo is missing


                # Title
                title_p = doc.add_paragraph()
                title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_run = title_p.add_run("LAPORAN HARIAN KEJADIAN BENCANA, WABAK, KECEMASAN, KRISIS (BWKK)\nPUSAT KESIAPSIAGAAN DAN TINDAKCEPAT KRISIS (CPRC)\nJABATAN KESIHATAN NEGERI SELANGOR\n")
                title_run.bold = True
                title_run.font.name = 'Arial'
                title_run.font.size = Pt(12)


                # Sub-header Table (Tarikh, Masa, Minggu Epid)
                table_hdr = doc.add_table(rows=1, cols=3)
                table_hdr.style = 'Table Grid'
                row = table_hdr.rows.cells
                row.text = f"Tarikh : {tarikh_today_str}"
                row.text = "(Sehingga jam 10.00 pagi)"
                row.text = f"Minggu Epidemiologi : {epid_week}/{today.year}"
                for cell in row:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.name = 'Arial'


                # Section 1.0
                doc.add_paragraph()
                p1 = doc.add_paragraph()
                p1.add_run("1.0 Ringkasan Laporan Input Enotifikasi").bold = True
                doc.add_paragraph(f"1.1 Sejumlah {total_enotif} input notifikasi telah diterima pada {tarikh_yesterday_str} dengan pecahan mengikut penyakit seperti dalam jadual 1.")
                doc.add_paragraph("Jadual 1 : Senarai Input Enotifikasi").bold = True
                
                # Table 1
                t1 = doc.add_table(rows=1, cols=len(df_jadual_1.columns)+1)
                t1.style = 'Table Grid'
                hdr_cells = t1.rows.cells
                hdr_cells.text = "PENYAKIT"
                for i, col in enumerate(df_jadual_1.columns):
                    hdr_cells[i+1].text = str(col).upper()
                    set_cell_background(hdr_cells[i+1], "D9D9D9")
                set_cell_background(hdr_cells, "D9D9D9")


                for index, row_data in df_jadual_1.iterrows():
                    row_cells = t1.add_row().cells
                    row_cells.text = str(index)
                    for i, val in enumerate(row_data):
                        row_cells[i+1].text = str(int(val))


                # Section 2.0
                doc.add_paragraph()
                p2 = doc.add_paragraph()
                p2.add_run("2.0 Ringkasan Laporan Notifikasi Wabak").bold = True
                
                if total_outbreak == 0:
                    doc.add_paragraph(f"2.1 Tiada notifikasi wabak diterima pada {tarikh_yesterday_str}.")
                    doc.add_paragraph("Jadual 2 : Senarai Notifikasi Wabak").bold = True
                    # Generate empty table
                    t2 = doc.add_table(rows=2, cols=3)
                    t2.style = 'Table Grid'
                    hdr_cells = t2.rows.cells
                    for i, text in enumerate(["PENYAKIT", "HARIAN", "KUMULATIF"]):
                        hdr_cells[i].text = text
                        set_cell_background(hdr_cells[i], "D9D9D9")
                    for cell in t2.rows.cells:
                        cell.text = "0"
                else:
                    doc.add_paragraph(f"2.1 Sejumlah {total_outbreak} input notifikasi wabak telah diterima pada {tarikh_yesterday_str} dengan pecahan mengikut penyakit seperti dalam jadual 2.")
                    doc.add_paragraph("Jadual 2 : Senarai Notifikasi Wabak").bold = True
                    t2 = doc.add_table(rows=1, cols=3)
                    t2.style = 'Table Grid'
                    hdr_cells = t2.rows.cells
                    for i, text in enumerate(["PENYAKIT", "HARIAN", "KUMULATIF"]):
                        hdr_cells[i].text = text
                        set_cell_background(hdr_cells[i], "D9D9D9")
                    
                    for _, row_data in df_jadual_2.iterrows():
                        row_cells = t2.add_row().cells
                        row_cells.text = str(row_data['PENYAKIT'])
                        row_cells.text = str(int(row_data['HARIAN']))
                        row_cells.text = str(int(row_data['KUMULATIF']))


                # Section 3.0
                doc.add_paragraph()
                p3 = doc.add_paragraph()
                p3.add_run("3.0 Ringkasan Laporan Input Enotifikasi").bold = True 
                
                if df_jadual_3.empty:
                    doc.add_paragraph(f"3.1 Tiada notifikasi wabak vektor diterima pada {tarikh_yesterday_str}.")
                else:
                    doc.add_paragraph(f"3.1 Sejumlah xx input notifikasi wabak vektor telah diterima pada {tarikh_yesterday_str} dengan pecahan mengikut penyakit seperti dalam jadual 3.")
                    doc.add_paragraph("Jadual 3 : Senarai Notifikasi Wabak Vektor").bold = True
                    
                    t3 = doc.add_table(rows=1, cols=len(df_jadual_3.columns))
                    t3.style = 'Table Grid'
                    hdr_cells = t3.rows.cells
                    for i, col in enumerate(df_jadual_3.columns):
                        hdr_cells[i].text = str(col)
                        set_cell_background(hdr_cells[i], "D9D9D9")
                    for _, row_data in df_jadual_3.iterrows():
                        row_cells = t3.add_row().cells
                        for i, val in enumerate(row_data):
                            row_cells[i].text = str(val)


                # Section 4.0
                doc.add_paragraph()
                p4 = doc.add_paragraph()
                p4.add_run("4.0 Ringkasan Laporan Kejadian Insiden").bold = True
                
                if total_insiden_yesterday == 0:
                    doc.add_paragraph(f"4.1 Tiada Insiden yang dilaporkan pada {tarikh_yesterday_str}.")
                else:
                    doc.add_paragraph(f"4.1 Sejumlah {total_insiden_yesterday} input notifikasi Kejadian Insiden Bencana, Kecemasan dan Krisis (BKK) telah diterima pada {tarikh_yesterday_str} dengan pecahan mengikut penyakit seperti dalam jadual 4.")
                    doc.add_paragraph("Jadual 4 : Senarai Kejadian Insiden Bencana, Kecemasan dan Krisis (BKK)").bold = True
                    
                    t4 = doc.add_table(rows=1, cols=len(df_jadual_4.columns))
                    t4.style = 'Table Grid'
                    hdr_cells = t4.rows.cells
                    for i, col in enumerate(df_jadual_4.columns):
                        hdr_cells[i].text = str(col)
                        set_cell_background(hdr_cells[i], "D9D9D9")
                    for _, row_data in df_jadual_4.iterrows():
                        row_cells = t4.add_row().cells
                        for i, val in enumerate(row_data):
                            row_cells[i].text = str(val)


                # Footer
                doc.add_paragraph()
                footer_text = f"*Sumber : Sistem e-notifikasi, Laporan Wabak KKM dimuat turun pada ({tarikh_today_str} @ 10.00 am)"
                doc.add_paragraph(footer_text)
                
                doc.add_paragraph("Petugas\t\t:\nJawatan \t\t:\nKetua Petugas\t:\nJawatan\t\t:")


                # 5. PREPARE FILE FOR DOWNLOAD
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                
                st.success("Laporan Berjaya Dijana!")
                
                st.download_button(
                    label="Muat Turun Laporan (Word)",
                    data=buffer,
                    file_name=f"Laporan_Harian_BWKK_{today.strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            
            except Exception as e:
                st.error(f"Terdapat ralat semasa menjana laporan: {str(e)}")
